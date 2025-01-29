from typing import List, Optional
import fitz  # PyMuPDF
from pathlib import Path
import tiktoken
import openai
from src.models.document_models import DocumentChunk, DocumentMetadata
import os
import docx
import asyncio

class DocumentProcessor:
    def __init__(self):
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        self.chunk_size = 1000
        self.chunk_overlap = 200
        self.openai_client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

    async def process_document(self, file_path: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """Process a document and return chunks"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._process_pdf(file_path, metadata)
        elif file_ext == '.docx':
            return await self._process_docx(file_path, metadata)
        elif file_ext in ['.txt', '.md']:
            return await self._process_text(file_path, metadata)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    async def _process_pdf(self, file_path: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        chunks = []
        doc = fitz.open(file_path)
        metadata.total_pages = len(doc)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            page_chunks = self._create_chunks(text, metadata.title, page_num + 1)
            chunks.extend(page_chunks)
            
        return chunks
    
    async def _process_docx(self, file_path: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        chunks = []
        doc = docx.Document(file_path)
        metadata.total_pages = len(doc.paragraphs)  # Approximate pages
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        chunks = self._create_chunks(text, metadata.title, None)
        return chunks
    
    async def _process_text(self, file_path: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        chunks = self._create_chunks(text, metadata.title, None)
        return chunks
    
    def _create_chunks(self, text: str, source: str, page_num: Optional[int]) -> List[DocumentChunk]:
        """Split text into chunks with overlap"""
        chunks = []
        
        # Split text into chunks
        words = text.split()
        current_chunk = []
        current_size = 0
        chunk_number = 0
        
        for word in words:
            current_chunk.append(word)
            current_size += len(word) + 1  # +1 for space
            
            if current_size >= self.chunk_size:
                # Create chunk
                chunk_text = " ".join(current_chunk)
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    source=source,
                    page_number=page_num,
                    chunk_number=chunk_number
                ))
                
                # Keep overlap for next chunk
                overlap_words = current_chunk[-self.chunk_overlap:]
                current_chunk = overlap_words
                current_size = sum(len(word) + 1 for word in overlap_words)
                chunk_number += 1
        
        # Add final chunk if there's remaining text
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(DocumentChunk(
                content=chunk_text,
                source=source,
                page_number=page_num,
                chunk_number=chunk_number
            ))
        
        return chunks

    def extract_text_from_pdf(self, file_path: str) -> List[tuple[str, int]]:
        doc = fitz.open(file_path)
        text_pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_pages.append((page.get_text(), page_num + 1))
        return text_pages

    def create_chunks(self, text: str, source: str, page_number: Optional[int] = None) -> List[DocumentChunk]:
        tokens = self.tokenizer.encode(text)
        chunks = []
        
        for i in range(0, len(tokens), self.chunk_size - self.chunk_overlap):
            chunk_tokens = tokens[i:i + self.chunk_size]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            chunks.append(
                DocumentChunk(
                    content=chunk_text,
                    source=source,
                    page_number=page_number,
                    chunk_number=len(chunks) + 1
                )
            )
        return chunks

    async def create_embeddings(self, chunks: List[DocumentChunk]) -> List[List[float]]:
        embeddings = []
        for chunk in chunks:
            response = self.openai_client.embeddings.create(
                input=chunk.content,
                model="text-embedding-3-small"
            )
            embeddings.append(response.data[0].embedding)
        return embeddings 